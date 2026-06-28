import io
import sys
from pathlib import Path

import pytest
import fitz
from docx import Document as DocxDocument
from docx.shared import Inches
from PIL import Image

SCRIPTS_DIR = Path(__file__).resolve().parent.parent.parent / "write-paper-notes" / "scripts"
sys.path.insert(0, str(SCRIPTS_DIR))


def _create_simple_pdf(path):
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((72, 72), "Hello World - Paper Notes", fontsize=12)
    page.insert_text((72, 100), "This is a test paper document.", fontsize=10)
    doc.save(str(path))
    doc.close()


def _create_multi_page_pdf(path):
    doc = fitz.open()
    for i in range(3):
        page = doc.new_page()
        page.insert_text((72, 72), f"Section {i+1}: Paper content", fontsize=12)
    doc.save(str(path))
    doc.close()


def _create_images_pdf(path):
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((72, 72), "Paper with figures", fontsize=12)
    tmp = fitz.open()
    tp = tmp.new_page()
    tp.insert_text((50, 100), "Figure 1: Results chart", fontsize=24)
    pix = tp.get_pixmap(dpi=150)
    tmp.close()
    page.insert_image((72, 100, 372, 350), pixmap=pix)
    doc.save(str(path))
    doc.close()


def _create_tiny_image_pdf(path):
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((72, 72), "Tiny images test", fontsize=12)
    tmp = fitz.open()
    tp = tmp.new_page(width=1, height=36)
    pix = tp.get_pixmap(dpi=72)
    tmp.close()
    page.insert_image((72, 100, 73, 136), pixmap=pix)
    doc.save(str(path))
    doc.close()


def _create_vector_pdf(path):
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((72, 72), "Paper with vector diagrams", fontsize=12)

    page.draw_rect(fitz.Rect(100, 150, 250, 200), color=(0, 0, 1), width=1.5)
    page.insert_text((120, 175), "Step 1", fontsize=10)

    page.draw_line(fitz.Point(175, 200), fitz.Point(175, 240), color=(0, 0, 1))

    page.draw_rect(fitz.Rect(100, 240, 250, 290), color=(0, 0, 1), width=1.5)
    page.insert_text((120, 265), "Step 2", fontsize=10)

    page.draw_circle(fitz.Point(400, 200), 30, color=(1, 0, 0), width=1.5)
    page.insert_text((390, 200), "Y/N", fontsize=8)

    doc.save(str(path))
    doc.close()


def _create_simple_docx(path):
    doc = DocxDocument()
    doc.add_paragraph("Hello World - Paper Notes")
    doc.add_paragraph("This is a test paper document.")
    doc.save(str(path))


def _create_docx_with_headings(path):
    doc = DocxDocument()
    doc.add_heading("Section 1", level=1)
    doc.add_heading("Subsection A", level=2)
    doc.add_paragraph("Content under subsection A.")
    doc.add_heading("Section 2", level=1)
    doc.add_paragraph("Content under section 2.")
    doc.save(str(path))


def _create_docx_with_table(path):
    doc = DocxDocument()
    doc.add_heading("Results", level=1)
    table = doc.add_table(rows=3, cols=3)
    table.cell(0, 0).text = "Name"
    table.cell(0, 1).text = "Score"
    table.cell(0, 2).text = "Grade"
    table.cell(1, 0).text = "Alice"
    table.cell(1, 1).text = "95"
    table.cell(1, 2).text = "A"
    table.cell(2, 0).text = "Bob"
    table.cell(2, 1).text = "82"
    table.cell(2, 2).text = "B"
    doc.save(str(path))


def _create_docx_with_image(path):
    doc = DocxDocument()
    doc.add_heading("Figures", level=1)
    doc.add_paragraph("Below is a sample figure:")

    buf = io.BytesIO()
    Image.new("RGB", (100, 50), color=(200, 200, 200)).save(buf, format="PNG")
    buf.seek(0)
    doc.add_picture(buf, width=Inches(2))

    doc.add_paragraph("End of figures.")
    doc.save(str(path))


def _create_docx_empty(path):
    doc = DocxDocument()
    doc.save(str(path))


def _create_docx_with_bullets(path):
    doc = DocxDocument()
    doc.add_heading("Findings", level=1)
    doc.add_paragraph("First item", style="List Bullet")
    doc.add_paragraph("Second item", style="List Bullet")
    doc.add_paragraph("Third item", style="List Bullet")
    doc.save(str(path))


def _create_docx_with_numbers(path):
    doc = DocxDocument()
    doc.add_heading("Steps", level=1)
    doc.add_paragraph("Step one", style="List Number")
    doc.add_paragraph("Step two", style="List Number")
    doc.add_paragraph("Step three", style="List Number")
    doc.save(str(path))


@pytest.fixture(scope="session")
def fixtures_dir(tmp_path_factory):
    d = tmp_path_factory.mktemp("fixtures")
    _create_simple_pdf(d / "simple.pdf")
    _create_multi_page_pdf(d / "multi_page.pdf")
    _create_images_pdf(d / "images.pdf")
    _create_vector_pdf(d / "vector.pdf")
    _create_tiny_image_pdf(d / "tiny.pdf")
    _create_simple_docx(d / "simple.docx")
    _create_docx_with_headings(d / "headings.docx")
    _create_docx_with_table(d / "table.docx")
    _create_docx_with_image(d / "image.docx")
    _create_docx_empty(d / "empty.docx")
    _create_docx_with_bullets(d / "bullets.docx")
    _create_docx_with_numbers(d / "numbers.docx")
    return d


@pytest.fixture
def simple_pdf(fixtures_dir):
    return fixtures_dir / "simple.pdf"


@pytest.fixture
def multi_page_pdf(fixtures_dir):
    return fixtures_dir / "multi_page.pdf"


@pytest.fixture
def images_pdf(fixtures_dir):
    return fixtures_dir / "images.pdf"


@pytest.fixture
def vector_pdf(fixtures_dir):
    return fixtures_dir / "vector.pdf"


@pytest.fixture
def tiny_image_pdf(fixtures_dir):
    return fixtures_dir / "tiny.pdf"


@pytest.fixture
def output_dir(tmp_path):
    d = tmp_path / "output"
    d.mkdir()
    return d


@pytest.fixture
def simple_docx(fixtures_dir):
    return fixtures_dir / "simple.docx"


@pytest.fixture
def headings_docx(fixtures_dir):
    return fixtures_dir / "headings.docx"


@pytest.fixture
def table_docx(fixtures_dir):
    return fixtures_dir / "table.docx"


@pytest.fixture
def image_docx(fixtures_dir):
    return fixtures_dir / "image.docx"


@pytest.fixture
def empty_docx(fixtures_dir):
    return fixtures_dir / "empty.docx"


@pytest.fixture
def bullets_docx(fixtures_dir):
    return fixtures_dir / "bullets.docx"


@pytest.fixture
def numbers_docx(fixtures_dir):
    return fixtures_dir / "numbers.docx"
